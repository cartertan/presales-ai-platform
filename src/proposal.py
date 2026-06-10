"""Proposal Word document generator for the Presales AI Platform.

Builds a 17-section draft proposal .docx from RFP analysis and RAG knowledge.
Each AI-generated section makes exactly one Ollama API call and parses the
response via XML markers.
"""

import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from dotenv import load_dotenv

from rag.retriever import format_context, search_industries, search_products, search_proposals

load_dotenv()

logger = logging.getLogger(__name__)

_WRITING_MODEL = os.getenv("OLLAMA_WRITING_MODEL", "qwen3.6:27b")
_COMPLIANCE_MODEL = os.getenv("OLLAMA_COMPLIANCE_MODEL", "granite4.1:30b")
_ANALYSIS_MODEL = os.getenv("OLLAMA_ANALYSIS_MODEL", "deepseek-r1:32b")
_FAST_MODEL = os.getenv("OLLAMA_FAST_MODEL", "gemma4:e4b")

_DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_GREY = RGBColor(0x60, 0x60, 0x60)
_TABLE_HEADER_HEX = "1F3864"

_SYS_WRITING = (
    "You are a professional presales proposal writer for Nexus Group, specialising in "
    "PKI, digital identity, and cybersecurity. Write formal, concise, enterprise-quality content. "
    "Always wrap your response in the specified XML tags."
)
_SYS_COMPLIANCE = (
    "You are a technical compliance specialist for Nexus Group PKI and digital identity products. "
    "Write precise, technically accurate content. Always wrap your response in the specified XML tags."
)
_SYS_ANALYSIS = (
    "You are a senior solution architect and project delivery expert for Nexus Group. "
    "Produce structured, analytical content from RFP context. "
    "Always wrap your response in the specified XML tags."
)


# ── Ollama and parsing helpers ────────────────────────────────────────────────

def _call_ollama(
    model: str,
    prompt: str,
    base_url: str,
    system: str = "",
    timeout: int = 300,
    num_predict: int = 4096,
) -> str:
    """Call the Ollama generate API and return the raw response text.

    Args:
        model: Ollama model name (e.g. qwen3.6:27b).
        prompt: User prompt text.
        base_url: Ollama API base URL.
        system: Optional system prompt string.
        timeout: HTTP request timeout in seconds.
        num_predict: Maximum tokens to generate.

    Returns:
        Raw model response text, stripped of leading/trailing whitespace.

    Raises:
        requests.exceptions.RequestException: On network or HTTP errors.
    """
    payload: dict[str, Any] = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {"temperature": 0.3, "num_predict": num_predict},
    }
    if system:
        payload["system"] = system
    logger.info("Calling Ollama model '%s' for proposal section", model)
    response = requests.post(
        f"{base_url}/api/generate", json=payload, timeout=timeout
    )
    response.raise_for_status()
    return response.json().get("response", "").strip()


def _extract_xml(text: str, tag: str) -> str:
    """Extract the content between <tag> and </tag> markers.

    Falls back to the full text if the tag is not present.

    Args:
        text: Raw model response text.
        tag: XML tag name without angle brackets.

    Returns:
        Extracted and stripped content string.
    """
    match = re.search(f"<{tag}>(.*?)</{tag}>", text, re.DOTALL)
    if match:
        return match.group(1).strip()
    return text.strip()


def _safe_rag(func: Any, *args: Any, **kwargs: Any) -> list:
    """Call a RAG search function and return [] if the knowledge base is unavailable.

    Args:
        func: RAG search function (e.g. search_products).
        *args: Positional arguments forwarded to func.
        **kwargs: Keyword arguments forwarded to func.

    Returns:
        Search result list, or empty list on any failure.
    """
    try:
        return func(*args, **kwargs)
    except Exception as exc:
        logger.warning("RAG search unavailable, using empty context: %s", exc)
        return []


def _extract_json_list(text: str) -> list[dict[str, Any]] | None:
    """Parse a JSON array from text using multiple strategies.

    Tries: direct parse → strip markdown fences → find outermost brackets.

    Args:
        text: Text containing a JSON array.

    Returns:
        Parsed list of dicts, or None if all attempts fail.
    """
    for candidate in [text, re.sub(r"```(?:json)?", "", text).strip()]:
        try:
            result = json.loads(candidate)
            if isinstance(result, list):
                return result
        except json.JSONDecodeError:
            pass
    start, end = text.find("["), text.rfind("]")
    if start != -1 and end > start:
        try:
            result = json.loads(text[start : end + 1])
            if isinstance(result, list):
                return result
        except json.JSONDecodeError:
            pass
    return None


def _slim_analysis(analysis: dict[str, Any]) -> dict[str, Any]:
    """Return a trimmed copy of the analysis dict to reduce prompt context size.

    Caps: executive_summary at 300 chars, requirements at top 10 with each
    requirement text truncated to 150 chars, win_themes and risk_areas at top 3.
    All other keys are passed through unchanged.

    Args:
        analysis: Full RFP analysis dict from analyzer.analyze_rfp().

    Returns:
        Shallow-copied dict with trimmed fields.
    """
    slim = dict(analysis)
    slim["executive_summary"] = analysis.get("executive_summary", "")[:300]
    slim["win_themes"] = analysis.get("win_themes", [])[:3]
    slim["risk_areas"] = analysis.get("risk_areas", [])[:3]
    slim_reqs = []
    for r in analysis.get("requirements", [])[:10]:
        slim_r = dict(r)
        slim_r["requirement"] = r.get("requirement", "")[:150]
        slim_reqs.append(slim_r)
    slim["requirements"] = slim_reqs
    return slim


# ── OOXML helpers ─────────────────────────────────────────────────────────────

def _shade_cell(cell: Any, hex_color: str) -> None:
    """Set the background fill of a table cell via OOXML.

    Args:
        cell: python-docx TableCell object.
        hex_color: Six-character hex colour string without '#', e.g. "1F3864".
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _append_field_code(para: Any, field_instr: str) -> None:
    """Append a Word field-code run (PAGE, NUMPAGES, etc.) to a paragraph.

    Args:
        para: python-docx Paragraph to append to.
        field_instr: Field instruction string, e.g. "PAGE".
    """
    run = para.add_run()
    run.font.size = Pt(9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    instr_elem = OxmlElement("w:instrText")
    instr_elem.set(qn("xml:space"), "preserve")
    instr_elem.text = f" {field_instr} "
    run._r.append(instr_elem)
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)
    end_elem = OxmlElement("w:fldChar")
    end_elem.set(qn("w:fldCharType"), "end")
    run._r.append(end_elem)


# ── Document structure helpers ────────────────────────────────────────────────

def _configure_document(doc: Document) -> None:
    """Set 1-inch page margins, Calibri 11pt body font, and the footer.

    Args:
        doc: The python-docx Document to configure.
    """
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Footer: Confidential — Nexus | Page X
    footer_section = doc.sections[0]
    footer = footer_section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = para.add_run("Confidential — Nexus  |  Page ")
    conf_run.font.size = Pt(9)
    conf_run.font.color.rgb = _GREY
    _append_field_code(para, "PAGE")


def _add_h1(doc: Document, text: str, page_break: bool = True) -> Any:
    """Add a Heading 1 paragraph with optional preceding page break.

    Args:
        doc: Document to append to.
        text: Heading text.
        page_break: Insert a page break before the heading (default True).

    Returns:
        The created Paragraph.
    """
    if page_break:
        doc.add_page_break()
    para = doc.add_paragraph(style="Heading 1")
    run = para.add_run(text)
    run.font.color.rgb = _DARK_BLUE
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    return para


def _add_h2(doc: Document, text: str) -> Any:
    """Add a Heading 2 paragraph.

    Args:
        doc: Document to append to.
        text: Heading text.

    Returns:
        The created Paragraph.
    """
    para = doc.add_paragraph(style="Heading 2")
    run = para.add_run(text)
    run.font.color.rgb = _DARK_BLUE
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Calibri"
    return para


def _body(doc: Document, text: str) -> Any:
    """Add a styled body paragraph.

    Args:
        doc: Document to append to.
        text: Paragraph text.

    Returns:
        The created Paragraph.
    """
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return para


def _add_paras(doc: Document, text: str) -> None:
    """Split text on blank lines and add each chunk as a body paragraph.

    Args:
        doc: Document to append to.
        text: Multi-paragraph text block.
    """
    for chunk in re.split(r"\n{2,}", text):
        chunk = chunk.strip()
        if chunk:
            _body(doc, chunk)


def _add_numbered_list(doc: Document, items: list[str]) -> None:
    """Render items as List Number style paragraphs.

    Args:
        doc: Document to append to.
        items: Text lines; leading numeric markers are stripped.
    """
    for item in items:
        clean = re.sub(r"^\d+[\.\)]\s*", "", item.strip())
        if clean:
            doc.add_paragraph(clean, style="List Number")


def _create_styled_table(doc: Document, headers: list[str]) -> Any:
    """Create a bordered table with a dark blue header row.

    Args:
        doc: Document to append to.
        headers: Column header labels.

    Returns:
        The created Table with header row populated.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for cell, header in zip(hdr_row.cells, headers):
        _shade_cell(cell, _TABLE_HEADER_HEX)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(header)
        run.font.color.rgb = _WHITE
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    return table


def _add_table_row(table: Any, values: list[str]) -> None:
    """Append a data row to a table.

    Args:
        table: python-docx Table to append to.
        values: Cell values for the new row.
    """
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        cell.text = str(value)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)


def _add_hr(doc: Document) -> None:
    """Add a horizontal rule via a paragraph bottom border.

    Args:
        doc: Document to append to.
    """
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3864")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _bold_sub(doc: Document, text: str) -> None:
    """Add a bold dark-blue subheading paragraph.

    Args:
        doc: Document to append to.
        text: Subheading text.
    """
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    run.font.color.rgb = _DARK_BLUE


# ── Section 1: Cover Page ─────────────────────────────────────────────────────

def _sec1_cover(
    doc: Document,
    analysis: dict[str, Any],
    selected_products: list[dict[str, Any]],
) -> None:
    """Build the cover page. No AI generation required.

    Args:
        doc: Document to append to.
        analysis: RFP analysis dict (uses _rfp_filename key).
        selected_products: Selected product dicts.
    """
    rfp_filename = analysis.get("_rfp_filename", "rfp")
    customer_name = Path(rfp_filename).stem.replace("_", " ").replace("-", " ").title()
    product_names = ", ".join(p.get("name", "") for p in selected_products)
    today_str = datetime.now().strftime("%d %B %Y")

    for _ in range(7):
        doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = title_para.add_run("Technical Proposal")
    t.font.color.rgb = _DARK_BLUE
    t.font.size = Pt(28)
    t.bold = True
    t.font.name = "Calibri"

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub_para.add_run(product_names or "Nexus PKI Solution")
    s.font.color.rgb = _DARK_BLUE
    s.font.size = Pt(14)
    s.font.name = "Calibri"

    doc.add_paragraph()

    for label, value in [
        ("Prepared for", customer_name),
        ("Prepared by", "Nexus Group"),
        ("Date", today_str),
        ("Version", "1.0 DRAFT"),
        ("Classification", "Confidential"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}:  {value}")
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    doc.add_page_break()


# ── Section 2: Executive Summary ──────────────────────────────────────────────

def _sec2_exec_summary(
    doc: Document,
    analysis: dict[str, Any],
    selected_products: list[dict[str, Any]],
    config: dict[str, Any],
) -> None:
    """Build Section 2: Executive Summary.

    RAG: product names (n=3) + industry vertical (n=2).
    Model: writing_model. XML tag: exec_summary.

    Args:
        doc: Document to append to.
        analysis: RFP analysis dict.
        selected_products: Selected product dicts.
        config: Runtime config dict.
    """
    _add_h1(doc, "1. Executive Summary", page_break=False)
    base_url = config["ollama_base_url"]
    model = config["writing_model"]
    slim = _slim_analysis(analysis)
    product_names = ", ".join(p.get("name", "") for p in selected_products)
    industry = analysis.get("industry_vertical", "")

    rag_results = (
        _safe_rag(search_products, product_names, n_results=3, base_url=base_url)
        + _safe_rag(search_industries, industry, n_results=2, base_url=base_url)
    )
    rag_ctx = format_context(rag_results, max_chars=1500)

    context = (
        f"EXECUTIVE SUMMARY FROM RFP ANALYSIS:\n{slim['executive_summary']}\n\n"
        f"WIN THEMES:\n{json.dumps(slim['win_themes'])}\n\n"
        f"SELECTED NEXUS PRODUCTS: {product_names}\n\n"
        f"KNOWLEDGE BASE CONTEXT:\n{rag_ctx}"
    )
    prompt = (
        "Write an executive summary for a Nexus Group technical proposal. "
        "Write exactly 3 professional paragraphs: "
        "(1) Acknowledge the customer's challenge and strategic requirements, "
        "(2) Describe how the selected Nexus products address those requirements directly, "
        "(3) Affirm Nexus Group's credentials, track record, and commitment to delivery. "
        "Wrap your full response in <exec_summary> tags.\n\nContext:\n" + context
    )

    raw = _call_ollama(model, prompt, base_url, system=_SYS_WRITING, num_predict=600)
    _add_paras(doc, _extract_xml(raw, "exec_summary"))
    doc.add_paragraph()


# ── Section 3: Understanding of Customer Requirements ─────────────────────────

def _sec3_understanding(
    doc: Document,
    analysis: dict[str, Any],
    config: dict[str, Any],
) -> None:
    """Build Section 3: Understanding of Customer Requirements.

    Generates 2 paragraphs of narrative, then a requirements table (max 15 rows).
    Model: writing_model. XML tag: understanding.

    Args:
        doc: Document to append to.
        analysis: RFP analysis dict.
        config: Runtime config dict.
    """
    _add_h1(doc, "2. Understanding of Customer Requirements")
    base_url = config["ollama_base_url"]
    model = config["writing_model"]
    slim = _slim_analysis(analysis)

    mandatory_top10 = [
        r for r in slim["requirements"] if str(r.get("type", "")).upper() == "MANDATORY"
    ][:10]

    context = (
        f"CUSTOMER OBJECTIVES:\n{json.dumps(analysis.get('customer_objectives', []))}\n\n"
        f"TOP MANDATORY REQUIREMENTS:\n"
        + "\n".join(
            f"- {r.get('id', '')}: {r.get('requirement', '')}"
            for r in mandatory_top10
        )
    )
    prompt = (
        "Write two professional paragraphs demonstrating Nexus Group's understanding of "
        "the customer's requirements. The first paragraph should summarise the overall "
        "business and technical challenge. The second paragraph should acknowledge the "
        "key mandatory requirements and our readiness to address them. "
        "Wrap your response in <understanding> tags.\n\nContext:\n" + context
    )

    raw = _call_ollama(model, prompt, base_url, system=_SYS_WRITING, num_predict=500)
    _add_paras(doc, _extract_xml(raw, "understanding"))
    doc.add_paragraph()

    # Requirements table — populated directly from analysis, no AI
    table_reqs = analysis.get("requirements", [])[:15]
    if table_reqs:
        table = _create_styled_table(
            doc, ["Req ID", "Requirement", "Type", "Priority"]
        )
        for req in table_reqs:
            _add_table_row(
                table,
                [
                    req.get("id", ""),
                    req.get("requirement", "")[:120],
                    req.get("type", ""),
                    req.get("scoring_weight") or "—",
                ],
            )
    doc.add_paragraph()


# ── Section 4: Proposed Solution Overview ─────────────────────────────────────

def _sec4_solution_overview(
    doc: Document,
    analysis: dict[str, Any],
    selected_products: list[dict[str, Any]],
    config: dict[str, Any],
) -> None:
    """Build Section 4: Proposed Solution Overview.

    RAG: product names (n=5). One overview paragraph plus one per product.
    Model: writing_model. XML tag: solution_overview.

    Args:
        doc: Document to append to.
        analysis: RFP analysis dict.
        selected_products: Selected product dicts.
        config: Runtime config dict.
    """
    _add_h1(doc, "3. Proposed Solution Overview")
    base_url = config["ollama_base_url"]
    model = config["writing_model"]
    slim = _slim_analysis(analysis)
    product_names = ", ".join(p.get("name", "") for p in selected_products)

    rag_results = _safe_rag(
        search_products, product_names, n_results=3, base_url=base_url
    )
    rag_ctx = format_context(rag_results, max_chars=1500)

    context = (
        f"SELECTED NEXUS PRODUCTS: {product_names}\n\n"
        f"CUSTOMER OBJECTIVES:\n{json.dumps(analysis.get('customer_objectives', []))}\n\n"
        f"PRODUCT KNOWLEDGE:\n{rag_ctx}"
    )
    prompt = (
        "Write a proposed solution overview for a Nexus Group technical proposal. "
        "First write one overview paragraph introducing the overall solution architecture and approach. "
        "Then write exactly one paragraph for each of the following products, describing how it "
        f"contributes to meeting the customer's requirements: {product_names}. "
        "Wrap your complete response in <solution_overview> tags.\n\nContext:\n" + context
    )

    raw = _call_ollama(
        model, prompt, base_url, system=_SYS_WRITING, num_predict=700,
    )
    _add_paras(doc, _extract_xml(raw, "solution_overview"))
    doc.add_paragraph()


# ── Section 5: Solution Architecture (BLANK) ──────────────────────────────────

def _sec5_architecture(doc: Document) -> None:
    """Build Section 5: Solution Architecture placeholder.

    Args:
        doc: Document to append to.
    """
    _add_h1(doc, "4. Solution Architecture")
    _body(
        doc,
        "The proposed solution architecture for this engagement is detailed below. "
        "[PLACEHOLDER: Insert architecture diagram]",
    )
    doc.add_paragraph()
    table = _create_styled_table(
        doc, ["Component", "Product", "Role", "Deployment Mode", "Notes"]
    )
    for _ in range(6):
        _add_table_row(table, ["", "", "", "", ""])
    doc.add_paragraph()
    _body(doc, "[TO BE COMPLETED BY SOLUTION ARCHITECT]")
    doc.add_paragraph()


# ── Section 6: Product Descriptions (RAG only) ────────────────────────────────

def _sec6_product_descriptions(
    doc: Document,
    selected_products: list[dict[str, Any]],
    config: dict[str, Any],
) -> None:
    """Build Section 6: Product Descriptions using RAG only — no AI generation.

    Retrieves product knowledge for each selected product and formats it as
    a bold subheading followed by paragraph content, separated by horizontal rules.

    Args:
        doc: Document to append to.
        selected_products: Selected product dicts.
        config: Runtime config dict.
    """
    _add_h1(doc, "5. Product Descriptions")
    base_url = config["ollama_base_url"]

    for i, product in enumerate(selected_products):
        product_name = product.get("name", "")
        results = _safe_rag(
            search_products, product_name, n_results=3, base_url=base_url
        )

        _bold_sub(doc, product_name)

        if results:
            for result in results:
                text = result.get("text", "").strip()
                if text:
                    _add_paras(doc, text[:800])
        else:
            _body(
                doc,
                f"{product_name}: {product.get('description', 'See Nexus Group product documentation.')}",
            )

        if i < len(selected_products) - 1:
            _add_hr(doc)

    doc.add_paragraph()


# ── Section 7: Compliance Statement ──────────────────────────────────────────

def _sec7_compliance(
    doc: Document,
    analysis: dict[str, Any],
    selected_products: list[dict[str, Any]],
    config: dict[str, Any],
) -> None:
    """Build Section 7: Compliance Statement.

    Asks compliance_model to map each requirement to a compliance status and Nexus
    product. Parses the JSON array from within <compliance> tags.

    Args:
        doc: Document to append to.
        analysis: RFP analysis dict.
        selected_products: Selected product dicts.
        config: Runtime config dict.
    """
    _add_h1(doc, "6. Compliance Statement")
    base_url = config["ollama_base_url"]
    model = config["compliance_model"]
    slim = _slim_analysis(analysis)
    product_names = ", ".join(p.get("name", "") for p in selected_products)

    reqs_summary = [
        {
            "id": r.get("id", ""),
            "requirement": r.get("requirement", ""),
            "type": r.get("type", ""),
            "nexus_solution": r.get("nexus_solution", ""),
        }
        for r in slim["requirements"]
    ]

    context = (
        f"NEXUS PRODUCTS IN THIS PROPOSAL: {product_names}\n\n"
        f"REQUIREMENTS:\n{json.dumps(reqs_summary, indent=2)}"
    )
    prompt = (
        "Assess the compliance status of each requirement against the Nexus products listed. "
        "Use only these status values: Fully Compliant, Partially Compliant, TBC.\n\n"
        "Return a JSON array inside <compliance> tags — no other text:\n"
        "<compliance>\n"
        "[\n"
        '  {"req_id": "RFP-001", "requirement": "short text", "type": "MANDATORY", '
        '"status": "Fully Compliant", "solution": "Product Name", "remarks": "brief note"}\n'
        "]\n"
        "</compliance>\n\n"
        "Context:\n" + context
    )

    raw = _call_ollama(
        model, prompt, base_url, system=_SYS_COMPLIANCE, num_predict=800
    )
    content = _extract_xml(raw, "compliance")
    rows = _extract_json_list(content)

    table = _create_styled_table(
        doc,
        ["Req ID", "Requirement", "Type", "Compliance Status", "Nexus Solution", "Remarks"],
    )

    if rows:
        for row in rows:
            _add_table_row(
                table,
                [
                    row.get("req_id", ""),
                    row.get("requirement", "")[:120],
                    row.get("type", ""),
                    row.get("status", "TBC"),
                    row.get("solution", ""),
                    row.get("remarks", ""),
                ],
            )
    else:
        logger.warning("Section 7: compliance JSON parse failed; inserting raw content")
        _add_table_row(table, ["—", content[:200], "—", "TBC", "—", "See analysis"])

    doc.add_paragraph()


# ── Section 8: Technical Sizing (BLANK) ───────────────────────────────────────

def _sec8_sizing(doc: Document) -> None:
    """Build Section 8: Technical Sizing placeholder.

    Args:
        doc: Document to append to.
    """
    _add_h1(doc, "7. Technical Sizing")
    _body(doc, "[PLACEHOLDER: Complete based on customer environment]")
    doc.add_paragraph()
    table = _create_styled_table(doc, ["Component", "Specification", "Qty", "Notes"])
    for component in [
        "Application Server",
        "Database Server",
        "HSM",
        "Load Balancer",
        "Storage",
        "Network",
        "Backup",
    ]:
        _add_table_row(table, [component, "", "", ""])
    doc.add_paragraph()
    _body(doc, "[TO BE COMPLETED: Insert customer environment sizing]")
    doc.add_paragraph()


# ── Section 9: Scope of Work ──────────────────────────────────────────────────

def _sec9_scope(
    doc: Document,
    analysis: dict[str, Any],
    selected_products: list[dict[str, Any]],
    config: dict[str, Any],
) -> None:
    """Build Section 9: Scope of Work.

    Generates a numbered deliverables list covering supply through go-live.
    Model: compliance_model. XML tag: scope.

    Args:
        doc: Document to append to.
        analysis: RFP analysis dict.
        selected_products: Selected product dicts.
        config: Runtime config dict.
    """
    _add_h1(doc, "8. Scope of Work")
    base_url = config["ollama_base_url"]
    model = config["compliance_model"]
    slim = _slim_analysis(analysis)
    product_names = ", ".join(p.get("name", "") for p in selected_products)

    mandatory_reqs = [
        r.get("requirement", "")[:100]
        for r in slim["requirements"]
        if str(r.get("type", "")).upper() == "MANDATORY"
    ][:10]

    context = (
        f"SELECTED NEXUS PRODUCTS: {product_names}\n\n"
        f"MANDATORY REQUIREMENTS:\n"
        + "\n".join(f"- {r}" for r in mandatory_reqs)
    )
    prompt = (
        "Write a numbered scope of work for a Nexus Group project. "
        "Include deliverables covering these areas in order: "
        "Supply (hardware/software licences), Installation, Configuration, "
        "Integration, Testing, Training, and Go-live support. "
        "Write one numbered item per deliverable. Be specific to the products and requirements. "
        "Wrap your full numbered list in <scope> tags.\n\nContext:\n" + context
    )

    raw = _call_ollama(
        model, prompt, base_url, system=_SYS_COMPLIANCE, num_predict=500
    )
    content = _extract_xml(raw, "scope")
    items = [ln.strip() for ln in content.splitlines() if ln.strip()]
    _add_numbered_list(doc, items)
    doc.add_paragraph()


# ── Section 10: Assumptions and Dependencies ──────────────────────────────────

def _sec10_assumptions(
    doc: Document,
    analysis: dict[str, Any],
    config: dict[str, Any],
) -> None:
    """Build Section 10: Assumptions and Dependencies.

    Model: analysis_model. XML tag: assumptions.

    Args:
        doc: Document to append to.
        analysis: RFP analysis dict.
        config: Runtime config dict.
    """
    _add_h1(doc, "9. Assumptions and Dependencies")
    base_url = config["ollama_base_url"]
    model = config["analysis_model"]
    slim = _slim_analysis(analysis)

    context = (
        f"RISK AREAS:\n{json.dumps(slim['risk_areas'])}\n\n"
        f"REQUIREMENTS SUMMARY:\n"
        + "\n".join(
            f"- {r.get('id', '')}: {r.get('requirement', '')}"
            for r in slim["requirements"]
        )
    )
    prompt = (
        "Based on the RFP context, generate two separate numbered lists. "
        "First list — ASSUMPTIONS: project assumptions Nexus Group is making. "
        "Second list — DEPENDENCIES: items the project depends on from the customer or third parties. "
        "Format exactly as:\n"
        "ASSUMPTIONS:\n1. ...\n2. ...\n\nDEPENDENCIES:\n1. ...\n2. ...\n"
        "Wrap the entire response in <assumptions> tags.\n\nContext:\n" + context
    )

    raw = _call_ollama(
        model, prompt, base_url, system=_SYS_ANALYSIS, num_predict=400
    )
    content = _extract_xml(raw, "assumptions")

    # Split into assumptions and dependencies blocks
    assumptions_block = ""
    dependencies_block = ""
    if "DEPENDENCIES:" in content.upper():
        parts = re.split(r"DEPENDENCIES\s*:", content, flags=re.IGNORECASE, maxsplit=1)
        assumptions_block = parts[0].strip()
        dependencies_block = parts[1].strip() if len(parts) > 1 else ""
        assumptions_block = re.sub(
            r"^ASSUMPTIONS\s*:", "", assumptions_block, flags=re.IGNORECASE
        ).strip()
    else:
        assumptions_block = content

    _add_h2(doc, "Assumptions")
    items_a = [ln.strip() for ln in assumptions_block.splitlines() if ln.strip()]
    _add_numbered_list(doc, items_a)
    doc.add_paragraph()

    if dependencies_block:
        _add_h2(doc, "Dependencies")
        items_d = [ln.strip() for ln in dependencies_block.splitlines() if ln.strip()]
        _add_numbered_list(doc, items_d)
        doc.add_paragraph()


# ── Section 11: Implementation Approach ───────────────────────────────────────

def _sec11_implementation(
    doc: Document,
    config: dict[str, Any],
) -> None:
    """Build Section 11: Implementation Approach.

    Generates 4 phases (Initiation, Installation, Integration and Testing, Go-Live)
    each with objectives, key activities, deliverables, and indicative duration.
    Model: compliance_model. XML tag: implementation.

    Args:
        doc: Document to append to.
        config: Runtime config dict.
    """
    _add_h1(doc, "10. Implementation Approach")
    base_url = config["ollama_base_url"]
    model = config["compliance_model"]

    prompt = (
        "Write an implementation approach for a Nexus Group PKI/digital identity project "
        "with exactly 4 phases. For each phase provide: Objectives, Key Activities, "
        "Deliverables, and Indicative Duration. "
        "Use this format for each phase:\n"
        "PHASE N: [Phase Name]\n"
        "Objectives: ...\n"
        "Key Activities: ...\n"
        "Deliverables: ...\n"
        "Indicative Duration: ...\n\n"
        "Phases must be: Phase 1: Initiation, Phase 2: Installation, "
        "Phase 3: Integration and Testing, Phase 4: Go-Live and Handover. "
        "Wrap your full response in <implementation> tags."
    )

    raw = _call_ollama(
        model, prompt, base_url, system=_SYS_COMPLIANCE, num_predict=600
    )
    content = _extract_xml(raw, "implementation")

    # Split on phase headings and render each as H2 + body
    phase_blocks = re.split(r"(?=PHASE\s+\d+\s*:)", content, flags=re.IGNORECASE)
    for block in phase_blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.splitlines()
        heading_line = lines[0].strip()
        body_text = "\n".join(lines[1:]).strip()
        # Convert "PHASE N: Name" → H2 heading
        heading = re.sub(r"^PHASE\s+\d+\s*:\s*", "", heading_line, flags=re.IGNORECASE).strip()
        _add_h2(doc, heading or heading_line)
        _add_paras(doc, body_text)

    doc.add_paragraph()


# ── Sections 12–14: Project Management, Support Model, Training (combined) ────

def _sec12_14_combined(
    doc: Document,
    config: dict[str, Any],
) -> None:
    """Build Sections 12–14 with a single Ollama call.

    Generates Project Management, Support Model, and Training content in one
    API call using three XML section markers, then renders each section into
    the document. Reduces 3 separate model calls to 1.
    Model: compliance_model. XML tags: project_mgmt, support, training.

    Args:
        doc: Document to append to.
        config: Runtime config dict.
    """
    base_url = config["ollama_base_url"]
    model = config["compliance_model"]

    prompt = (
        "Generate three sections for a Nexus Group PKI/digital identity proposal. "
        "Return XML only — no preamble, no explanation.\n\n"
        "Section 1 — Project Management. Write 4 short subsections (1-2 sentences each):\n"
        "1. Governance Structure, 2. RACI Summary, 3. Communication Plan, 4. Risk Management.\n"
        "Wrap in <project_mgmt> tags.\n\n"
        "Section 2 — Support Model. Return a JSON array of exactly 3 support tiers.\n"
        'Format: [{"tier": "...", "coverage": "...", "response_time": "...", '
        '"resolution_target": "...", "scope": "..."}]\n'
        "Wrap in <support> tags.\n\n"
        "Section 3 — Training. Return a JSON array of exactly 3 training packages.\n"
        'Format: [{"package": "...", "audience": "...", "duration": "...", "content": "..."}]\n'
        "Wrap in <training> tags."
    )

    raw = _call_ollama(
        model, prompt, base_url, system=_SYS_COMPLIANCE, num_predict=1000
    )

    # ── Section 12: Project Management ───────────────────────────────────────
    _add_h1(doc, "11. Project Management")
    pm_content = _extract_xml(raw, "project_mgmt")
    subsection_blocks = re.split(
        r"(?=\d+\.\s+(?:Governance|RACI|Communication|Risk))",
        pm_content,
        flags=re.IGNORECASE,
    )
    if len(subsection_blocks) <= 1:
        _add_paras(doc, pm_content)
    else:
        for block in subsection_blocks:
            block = block.strip()
            if not block:
                continue
            lines = block.splitlines()
            heading_match = re.match(r"^\d+\.\s+(.+)", lines[0].strip())
            if heading_match:
                _add_h2(doc, heading_match.group(1).strip())
                _add_paras(doc, "\n".join(lines[1:]).strip())
            else:
                _add_paras(doc, block)
    doc.add_paragraph()

    # ── Section 13: Support Model ─────────────────────────────────────────────
    _add_h1(doc, "12. Support Model")
    support_content = _extract_xml(raw, "support")
    tiers = _extract_json_list(support_content)
    table_support = _create_styled_table(
        doc, ["Tier", "Coverage", "Response Time", "Resolution Target", "Scope"]
    )
    if tiers:
        for tier in tiers:
            _add_table_row(
                table_support,
                [
                    tier.get("tier", ""),
                    tier.get("coverage", ""),
                    tier.get("response_time", ""),
                    tier.get("resolution_target", ""),
                    tier.get("scope", ""),
                ],
            )
    else:
        logger.warning("Combined call: support JSON parse failed; inserting raw content")
        _add_paras(doc, support_content)
    doc.add_paragraph()

    # ── Section 14: Training ──────────────────────────────────────────────────
    _add_h1(doc, "13. Training")
    training_content = _extract_xml(raw, "training")
    packages = _extract_json_list(training_content)
    table_training = _create_styled_table(
        doc, ["Training Package", "Audience", "Duration", "Content"]
    )
    if packages:
        for pkg in packages:
            _add_table_row(
                table_training,
                [
                    pkg.get("package", ""),
                    pkg.get("audience", ""),
                    pkg.get("duration", ""),
                    pkg.get("content", ""),
                ],
            )
    else:
        logger.warning("Combined call: training JSON parse failed; inserting raw content")
        _add_paras(doc, training_content)
    doc.add_paragraph()


# ── Section 15: Project Timeline (BLANK) ──────────────────────────────────────

def _sec15_timeline(doc: Document) -> None:
    """Build Section 15: Project Timeline placeholder.

    Args:
        doc: Document to append to.
    """
    _add_h1(doc, "14. Project Timeline")
    _body(doc, "[PLACEHOLDER: Insert project timeline]")
    doc.add_paragraph()
    table = _create_styled_table(
        doc, ["Phase", "Activity", "Duration", "Start", "End", "Dependencies"]
    )
    for _ in range(8):
        _add_table_row(table, ["", "", "", "", "", ""])
    doc.add_paragraph()
    _body(doc, "[TO BE COMPLETED: Confirm timeline with customer]")
    doc.add_paragraph()


# ── Section 16: Commercials (BLANK) ───────────────────────────────────────────

def _sec16_commercials(doc: Document) -> None:
    """Build Section 16: Commercials placeholder.

    Args:
        doc: Document to append to.
    """
    _add_h1(doc, "15. Commercials")
    _body(doc, "[PLACEHOLDER: Insert commercial pricing]")
    doc.add_paragraph()
    table = _create_styled_table(
        doc, ["Item", "Product / Service", "Qty", "Unit Price", "Total", "Notes"]
    )
    for _ in range(10):
        _add_table_row(table, ["", "", "", "", "", ""])
    doc.add_paragraph()
    _body(doc, "[TO BE COMPLETED: Insert pricing from commercial team]")
    doc.add_paragraph()


# ── Section 17: Company Profile and References ────────────────────────────────

def _sec17_company_profile(
    doc: Document,
    analysis: dict[str, Any],
    config: dict[str, Any],
) -> None:
    """Build Section 17: Company Profile and References.

    RAG: past proposals (n=3) + industry vertical (n=2).
    Model: writing_model. XML tag: company_profile.

    Args:
        doc: Document to append to.
        analysis: RFP analysis dict.
        config: Runtime config dict.
    """
    _add_h1(doc, "16. Company Profile and References")
    base_url = config["ollama_base_url"]
    model = config["writing_model"]
    industry = analysis.get("industry_vertical", "")

    rag_results = (
        _safe_rag(search_proposals, "Nexus PKI deployment", n_results=2, base_url=base_url)
        + _safe_rag(search_industries, industry, n_results=2, base_url=base_url)
    )
    rag_ctx = format_context(rag_results, max_chars=1500)

    context = f"INDUSTRY CONTEXT:\n{rag_ctx}"
    prompt = (
        "Write one professional paragraph introducing Nexus Group as a company — "
        "covering its history, specialisation in PKI and digital identity, geographic reach, "
        "and key differentiators. "
        "Then provide a JSON array of customer references inside <company_profile> tags. "
        "The paragraph comes first, then the JSON array. "
        "JSON format for references:\n"
        '[{"industry": "...", "solution": "...", "scale": "...", "outcome": "..."}]\n'
        "Extract reference data from the knowledge base context provided. "
        "If no specific references are available, create 3 representative examples "
        "consistent with Nexus Group's typical deployments. "
        "Wrap everything in <company_profile> tags.\n\nContext:\n" + context
    )

    raw = _call_ollama(
        model, prompt, base_url, system=_SYS_WRITING, num_predict=400
    )
    content = _extract_xml(raw, "company_profile")

    # Split paragraph from JSON array
    json_start = content.find("[")
    if json_start > 0:
        intro_text = content[:json_start].strip()
        json_text = content[json_start:].strip()
        _add_paras(doc, intro_text)
        doc.add_paragraph()
        refs = _extract_json_list(json_text)
    else:
        _add_paras(doc, content)
        refs = None

    if refs:
        table = _create_styled_table(
            doc, ["Industry", "Solution", "Scale", "Outcome"]
        )
        for ref in refs:
            _add_table_row(
                table,
                [
                    ref.get("industry", ""),
                    ref.get("solution", ""),
                    ref.get("scale", ""),
                    ref.get("outcome", ""),
                ],
            )
    doc.add_paragraph()


# ── Public API ────────────────────────────────────────────────────────────────

def generate_proposal(
    analysis: dict[str, Any],
    selected_products: list[dict[str, Any]],
    output_dir: str,
    config: dict[str, Any],
) -> str:
    """Generate a professional draft proposal Word document from RFP analysis.

    Builds a 17-section .docx containing cover page, AI-generated narrative
    sections, compliance matrix, BLANK placeholder sections, and RAG-sourced
    product/company content.

    Saved to: output_dir/proposals/{rfp_stem[:25]}_Proposal_Draft_{YYYYMMDD}.docx

    Args:
        analysis: Structured analysis dict from analyzer.analyze_rfp(), must
            include a '_rfp_filename' key (set by main.py).
        selected_products: List of selected product dicts from product_selector.
        output_dir: Root output directory; proposals/ subdirectory is created.
        config: Runtime configuration:
            - ollama_base_url (str): Ollama API base URL.
            - writing_model (str): Model for narrative sections (qwen3.6:27b).
            - compliance_model (str): Model for compliance/technical sections.
            - analysis_model (str): Model for reasoning-heavy sections.
            - fast_model (str): Fast model for simple template sections (gemma4:e4b).

    Returns:
        Absolute path to the saved .docx file.

    Raises:
        OSError: If the output directory cannot be created or file cannot be saved.
    """
    ollama_base_url = config.get("ollama_base_url", "http://localhost:11434")
    writing_model = config.get("writing_model", _WRITING_MODEL)
    compliance_model = config.get("compliance_model", _COMPLIANCE_MODEL)
    analysis_model = config.get("analysis_model", _ANALYSIS_MODEL)
    fast_model = config.get("fast_model", _FAST_MODEL)

    resolved_config: dict[str, Any] = {
        "ollama_base_url": ollama_base_url,
        "writing_model": writing_model,
        "compliance_model": compliance_model,
        "analysis_model": analysis_model,
        "fast_model": fast_model,
    }

    proposals_dir = Path(output_dir) / "proposals"
    try:
        proposals_dir.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        logger.error("Cannot create proposals directory '%s': %s", proposals_dir, exc)
        raise

    rfp_filename = analysis.get("_rfp_filename", "rfp")
    stem = Path(rfp_filename).stem[:25]
    date_str = datetime.now().strftime("%Y%m%d")
    filename = f"{stem}_Proposal_Draft_{date_str}.docx"
    file_path = proposals_dir / filename

    logger.info("Generating proposal draft: %s", file_path)

    doc = Document()
    _configure_document(doc)

    # Section 1 — Cover Page
    logger.info("Section 1: Cover Page")
    _sec1_cover(doc, analysis, selected_products)

    # Section 2 — Executive Summary
    logger.info("Section 2: Executive Summary [%s]", writing_model)
    _sec2_exec_summary(doc, analysis, selected_products, resolved_config)

    # Section 3 — Understanding of Customer Requirements
    logger.info("Section 3: Understanding of Customer Requirements [%s]", writing_model)
    _sec3_understanding(doc, analysis, resolved_config)

    # Section 4 — Proposed Solution Overview
    logger.info("Section 4: Proposed Solution Overview [%s]", writing_model)
    _sec4_solution_overview(doc, analysis, selected_products, resolved_config)

    # Section 5 — Solution Architecture (BLANK)
    logger.info("Section 5: Solution Architecture [BLANK]")
    _sec5_architecture(doc)

    # Section 6 — Product Descriptions (RAG only)
    logger.info("Section 6: Product Descriptions [RAG only]")
    _sec6_product_descriptions(doc, selected_products, resolved_config)

    # Section 7 — Compliance Statement
    logger.info("Section 7: Compliance Statement [%s]", compliance_model)
    _sec7_compliance(doc, analysis, selected_products, resolved_config)

    # Section 8 — Technical Sizing (BLANK)
    logger.info("Section 8: Technical Sizing [BLANK]")
    _sec8_sizing(doc)

    # Section 9 — Scope of Work
    logger.info("Section 9: Scope of Work [%s]", compliance_model)
    _sec9_scope(doc, analysis, selected_products, resolved_config)

    # Section 10 — Assumptions and Dependencies
    logger.info("Section 10: Assumptions and Dependencies [%s]", analysis_model)
    _sec10_assumptions(doc, analysis, resolved_config)

    # Section 11 — Implementation Approach
    logger.info("Section 11: Implementation Approach [%s]", compliance_model)
    _sec11_implementation(doc, resolved_config)

    # Sections 12–14 — Project Management, Support Model, Training (combined call)
    logger.info(
        "Sections 12–14: Project Management + Support + Training [%s, 1 combined call]",
        compliance_model,
    )
    _sec12_14_combined(doc, resolved_config)

    # Section 15 — Project Timeline (BLANK)
    logger.info("Section 15: Project Timeline [BLANK]")
    _sec15_timeline(doc)

    # Section 16 — Commercials (BLANK)
    logger.info("Section 16: Commercials [BLANK]")
    _sec16_commercials(doc)

    # Section 17 — Company Profile and References
    logger.info("Section 17: Company Profile and References [%s]", writing_model)
    _sec17_company_profile(doc, analysis, resolved_config)

    try:
        doc.save(str(file_path))
    except OSError as exc:
        logger.error("Failed to save proposal document '%s': %s", file_path, exc)
        raise

    resolved = str(file_path.resolve())
    logger.info("Proposal draft saved: %s", resolved)
    return resolved
