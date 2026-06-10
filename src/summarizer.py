"""Word document generator: creates a professional customer summary from RFP analysis."""

import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

SUMMARY_MODEL = os.getenv("OLLAMA_COMPLIANCE_MODEL", "granite4.1:30b")
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")

_DARK_BLUE = RGBColor(0x1F, 0x38, 0x64)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_GREY_TEXT = RGBColor(0x60, 0x60, 0x60)
_TABLE_HEADER_HEX = "1F3864"
_ALT_ROW_HEX = "EEF2F7"


# ── XML / OOXML helpers ───────────────────────────────────────────────────────

def _shade_cell(cell: Any, hex_color: str) -> None:
    """Set the background fill of a table cell via OOXML.

    Args:
        cell: python-docx TableCell object.
        hex_color: Six-character hex string without '#', e.g. "1F3864".
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _append_field_code(para: Any, field_instr: str) -> None:
    """Append a Word field-code run (PAGE, NUMPAGES, etc.) to a paragraph.

    Args:
        para: python-docx Paragraph to append to.
        field_instr: Field instruction, e.g. "PAGE" or "NUMPAGES".
    """
    run = para.add_run()
    run.font.size = Pt(9)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_elem = OxmlElement("w:instrText")
    instr_elem.set(qn("xml:space"), "preserve")
    instr_elem.text = f" {field_instr} "
    run._r.append(instr_elem)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)

    end_elem = OxmlElement("w:fldChar")
    end_elem.set(qn("w:fldCharType"), "end")
    run._r.append(end_elem)


def _add_page_numbers(doc: Document) -> None:
    """Insert 'Page X of Y' field codes in the document footer.

    Args:
        doc: The python-docx Document to modify.
    """
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pre = para.add_run("Page ")
    pre.font.size = Pt(9)
    _append_field_code(para, "PAGE")
    mid = para.add_run(" of ")
    mid.font.size = Pt(9)
    _append_field_code(para, "NUMPAGES")


# ── Document structure builders ───────────────────────────────────────────────

def _add_title_page(
    doc: Document, customer_name: str, project_title: str, display_date: str
) -> None:
    """Create a formatted title page with a trailing page break.

    Args:
        doc: Document to append to.
        customer_name: Organisation or customer name.
        project_title: RFP or project title.
        display_date: Human-readable date string, e.g. "8 June 2026".
    """
    for _ in range(7):
        doc.add_paragraph()

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = brand.add_run("NEXUS GROUP")
    br.font.color.rgb = _DARK_BLUE
    br.font.size = Pt(11)
    br.bold = True
    br.font.all_caps = True

    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_para.add_run(project_title)
    tr.font.color.rgb = _DARK_BLUE
    tr.font.size = Pt(26)
    tr.bold = True

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle_para.add_run("Customer Summary & RFP Analysis")
    sr.font.color.rgb = _DARK_BLUE
    sr.font.size = Pt(16)

    doc.add_paragraph()

    for label, value in [("Prepared for", customer_name), ("Date", display_date)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}:  {value}")
        r.font.size = Pt(12)

    doc.add_paragraph()

    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = conf_para.add_run("CONFIDENTIAL — FOR INTERNAL USE ONLY")
    cr.font.size = Pt(9)
    cr.font.color.rgb = _GREY_TEXT
    cr.italic = True

    doc.add_page_break()


def _add_heading(doc: Document, text: str, level: int = 1) -> Any:
    """Add a dark blue heading paragraph.

    Args:
        doc: Document to append to.
        text: Heading text.
        level: Heading level (1 or 2).

    Returns:
        The created Paragraph.
    """
    para = doc.add_paragraph(style=f"Heading {level}")
    run = para.add_run(text)
    run.font.color.rgb = _DARK_BLUE
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    return para


def _create_styled_table(doc: Document, headers: list[str]) -> Any:
    """Create a bordered table with a dark blue header row.

    Args:
        doc: Document to append to.
        headers: Column header labels.

    Returns:
        The created Table (header row already populated).
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for cell, header in zip(hdr_row.cells, headers):
        _shade_cell(cell, _TABLE_HEADER_HEX)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(header)
        run.font.color.rgb = _WHITE
        run.bold = True
        run.font.size = Pt(10)
    return table


def _add_table_row(
    table: Any, values: list[str], shading: str | None = None
) -> None:
    """Append a data row to a table.

    Args:
        table: python-docx Table to append to.
        values: Cell values for the new row.
        shading: Optional hex colour for the row background.
    """
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        cell.text = str(value)
        if shading:
            _shade_cell(cell, shading)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    """Render a list of strings as 'List Bullet' paragraphs.

    Args:
        doc: Document to append to.
        items: Text lines; leading bullet markers are stripped.
    """
    for item in items:
        clean = re.sub(r"^[\•\-\*\–]\s*", "", item.strip())
        if clean:
            doc.add_paragraph(clean, style="List Bullet")


def _add_numbered_list(doc: Document, items: list[str]) -> None:
    """Render a list of strings as 'List Number' paragraphs.

    Args:
        doc: Document to append to.
        items: Text lines; leading number markers are stripped.
    """
    for item in items:
        clean = re.sub(r"^\d+[\.\)]\s*", "", item.strip())
        if clean:
            doc.add_paragraph(clean, style="List Number")


# ── Context and model helpers ─────────────────────────────────────────────────

def _build_context(
    analysis: dict[str, Any], keys: list[str], max_chars: int = 6000
) -> str:
    """Concatenate selected analysis fields into a compact context string.

    Args:
        analysis: Full RFP analysis dict.
        keys: Keys to include.
        max_chars: Maximum character length before truncation.

    Returns:
        Formatted context string ready to pass to the model.
    """
    parts: list[str] = []
    for key in keys:
        value = analysis.get(key)
        if value:
            label = key.upper().replace("_", " ")
            parts.append(f"{label}:\n{json.dumps(value, indent=2, ensure_ascii=False)}")
    text = "\n\n".join(parts)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n[context truncated]"
    return text


def _extract_json_list(raw: str) -> list[dict[str, Any]] | None:
    """Attempt to parse a JSON array from raw model output.

    Tries direct parse, markdown-fence strip, then bracket search.

    Args:
        raw: Raw text potentially containing a JSON array.

    Returns:
        Parsed list, or None if all parse attempts fail.
    """
    for candidate in [
        raw,
        re.sub(r"```(?:json)?", "", raw).strip(),
    ]:
        try:
            result = json.loads(candidate)
            if isinstance(result, list):
                return result
        except json.JSONDecodeError:
            pass

    start, end = raw.find("["), raw.rfind("]")
    if start != -1 and end > start:
        try:
            result = json.loads(raw[start : end + 1])
            if isinstance(result, list):
                return result
        except json.JSONDecodeError:
            pass

    return None


def generate_section(prompt: str, context: str, model: str, base_url: str) -> str:
    """Generate text for one document section using Ollama.

    Retained for external callers. generate_summary() uses a single batched
    call instead and no longer calls this function.

    Args:
        prompt: Instruction describing what to generate for this section.
        context: Supporting context extracted from the RFP analysis.
        model: Ollama model name (e.g. "granite4.1:30b").
        base_url: Ollama API base URL.

    Returns:
        Generated text stripped of leading/trailing whitespace.

    Raises:
        requests.exceptions.RequestException: If the model call fails.
    """
    payload: dict[str, Any] = {
        "model": model,
        "prompt": f"{prompt}\n\nContext:\n{context}",
        "stream": False,
        "options": {"temperature": 0.3, "num_predict": 2048},
    }
    logger.info("Calling Ollama model '%s' for section generation", model)
    resp = requests.post(f"{base_url}/api/generate", json=payload, timeout=300)
    resp.raise_for_status()
    return resp.json().get("response", "").strip()


# ── XML section parser ────────────────────────────────────────────────────────

def _parse_xml_sections(raw: str) -> dict[str, str]:
    """Extract 8 named XML sections from the model response.

    Args:
        raw: Raw model output expected to contain XML-tagged sections.

    Returns:
        Dict mapping each tag name to its text content. Missing tags are
        replaced with a fixed fallback string.
    """
    fallbacks: dict[str, str] = {
        "section1_objectives": "Customer objectives to be confirmed with the customer.",
        "section2_use_cases": "Use cases to be identified during customer engagement.",
        "section3_technical": "Technical requirements summary pending review.",
        "section4_commercial": "Commercial requirements to be confirmed.",
        "section5_evaluation": "Evaluation criteria to be confirmed.",
        "section6_risks": "Risk assessment pending detailed review.",
        "section7_questions": "Clarification questions to be developed.",
        "section8_win_themes": "Win themes to be developed with the sales team.",
    }
    result: dict[str, str] = {}
    for tag, fallback in fallbacks.items():
        match = re.search(rf"<{tag}>(.*?)</{tag}>", raw, re.DOTALL | re.IGNORECASE)
        if match:
            result[tag] = match.group(1).strip()
            continue
        prefix_m = re.match(r"section\d+", tag)
        if prefix_m:
            partial = re.search(
                rf"<({re.escape(prefix_m.group(0))}[^>]*)>(.*?)</\1>",
                raw,
                re.DOTALL | re.IGNORECASE,
            )
            if partial:
                logger.warning("XML <%s>: exact tag missing, matched <%s>.", tag, partial.group(1))
                result[tag] = partial.group(2).strip()
                continue
        logger.warning("XML tag <%s> not found; using raw response.", tag)
        result[tag] = raw.strip() if raw else fallback
    return result


# ── Per-section builders ──────────────────────────────────────────────────────

def _section1_rfp_overview(
    doc: Document,
    analysis: dict[str, Any],
    selected_products: list[dict[str, Any]],
) -> None:
    """Add Section 1: RFP Overview as a two-column table."""
    _add_heading(doc, "1. RFP Overview")

    requirements = analysis.get("requirements", [])
    mandatory = sum(
        1 for r in requirements if str(r.get("type", "")).upper() == "MANDATORY"
    )
    optional = len(requirements) - mandatory

    rfp_name = analysis.get("_rfp_filename", "Unknown")
    customer_name = Path(rfp_name).stem.replace("_", " ").replace("-", " ").title()

    product_names = (
        ", ".join(p.get("name", "") for p in selected_products)
        if selected_products
        else "TBC — to be selected"
    )

    table = _create_styled_table(doc, ["Field", "Value"])
    rows = [
        ("Customer / Organisation", customer_name),
        ("Project Title", rfp_name),
        ("Industry Vertical", analysis.get("industry_vertical") or "Not specified"),
        ("Submission Deadline", analysis.get("submission_deadline") or "Not specified"),
        ("Budget Range", analysis.get("budget") or "Not specified"),
        (
            "Requirements",
            f"{len(requirements)} total  —  {mandatory} Mandatory, {optional} Optional",
        ),
        ("Proposed Nexus Products", product_names),
    ]
    for i, (field, value) in enumerate(rows):
        _add_table_row(table, [field, value], _ALT_ROW_HEX if i % 2 == 0 else None)

    doc.add_paragraph()


def _section2_customer_objectives(doc: Document, content: str) -> None:
    """Add Section 2: Customer Objectives as a bullet list.

    Args:
        doc: Document to append to.
        content: Pre-generated text; one objective per line.
    """
    _add_heading(doc, "2. Customer Objectives")
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    _add_bullet_list(doc, lines[:8])
    doc.add_paragraph()


def _section3_use_cases(doc: Document, content: str) -> None:
    """Add Section 3: Identified Use Cases as a numbered list.

    Args:
        doc: Document to append to.
        content: Pre-generated text; one use case per line.
    """
    _add_heading(doc, "3. Identified Use Cases")
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    _add_numbered_list(doc, lines[:10])
    doc.add_paragraph()


def _section4_technical_requirements(
    doc: Document, analysis: dict[str, Any], content: str
) -> None:
    """Add Section 4: Technical Requirements Summary.

    Attempts to parse content as a JSON array for a categorised table; falls
    back to a single-row table containing the plain text.

    Args:
        doc: Document to append to.
        analysis: Full RFP analysis dict (unused directly; kept for signature consistency).
        content: Pre-generated text for this section.
    """
    _add_heading(doc, "4. Technical Requirements Summary")
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    _add_bullet_list(doc, lines[:15])
    doc.add_paragraph()


def _section5_commercial_requirements(doc: Document, content: str) -> None:
    """Add Section 5: Commercial Requirements as a bullet list.

    Args:
        doc: Document to append to.
        content: Pre-generated text; one commercial requirement per line.
    """
    _add_heading(doc, "5. Commercial Requirements")
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    _add_bullet_list(doc, lines[:10])
    doc.add_paragraph()


def _section6_evaluation_criteria(
    doc: Document, analysis: dict[str, Any], content: str
) -> None:
    """Add Section 6: Evaluation Criteria as a table plus commentary paragraph.

    Table rows come from analysis.evaluation_criteria; the pre-generated
    content is added as a commentary paragraph below the table.

    Args:
        doc: Document to append to.
        analysis: Full RFP analysis dict (source of evaluation_criteria).
        content: Pre-generated evaluation commentary text.
    """
    _add_heading(doc, "6. Evaluation Criteria")

    criteria = analysis.get("evaluation_criteria", [])
    table = _create_styled_table(doc, ["Criterion", "Weight", "What It Means for Us"])

    if criteria:
        for i, crit in enumerate(criteria):
            name = crit.get("criterion", crit.get("name", ""))
            _add_table_row(
                table,
                [name, crit.get("weight", ""), ""],
                _ALT_ROW_HEX if i % 2 == 0 else None,
            )
    else:
        _add_table_row(table, ["Not specified", "—", "—"])

    if content:
        doc.add_paragraph(content[:800])

    doc.add_paragraph()


def _section7_risks_and_gaps(doc: Document, content: str) -> None:
    """Add Section 7: Key Risks and Gaps.

    Attempts to parse content as a JSON array for a colour-coded risk table;
    falls back to a bullet list when JSON parsing fails.

    Args:
        doc: Document to append to.
        content: Pre-generated risk/gap text.
    """
    _add_heading(doc, "7. Key Risks and Gaps")
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    _add_bullet_list(doc, lines[:12])
    doc.add_paragraph()


def _section8_clarification_questions(doc: Document, content: str) -> None:
    """Add Section 8: Clarification Questions as a numbered list.

    Args:
        doc: Document to append to.
        content: Pre-generated text; one question per line.
    """
    _add_heading(doc, "8. Clarification Questions")
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    _add_numbered_list(doc, lines[:20])
    doc.add_paragraph()


def _section9_win_themes(doc: Document, content: str) -> None:
    """Add Section 9: Win Themes as a bullet list.

    Args:
        doc: Document to append to.
        content: Pre-generated win-theme text; one theme per line.
    """
    _add_heading(doc, "9. Win Themes")
    lines = [ln.strip() for ln in content.splitlines() if ln.strip()]
    _add_bullet_list(doc, lines[:10])
    doc.add_paragraph()


# ── Public API ────────────────────────────────────────────────────────────────

def generate_summary(
    analysis: dict[str, Any], config: dict[str, Any], output_dir: str
) -> str:
    """Generate a professional customer summary Word document from RFP analysis.

    Makes a single Ollama API call to granite4.1:30b (OLLAMA_COMPLIANCE_MODEL)
    that produces all 8 content sections simultaneously via XML markers. The
    response is parsed and each section is written to the corresponding part of
    the Word document.

    File saved to: output_dir/summaries/{rfp_stem}_Summary_{YYYYMMDD}.docx

    Args:
        analysis: Structured dict from analyzer.analyze_rfp(), must include
            a '_rfp_filename' key (set by main.py after analysis).
        config: Runtime configuration:
            - base_url (str): Ollama API base URL.
            - selected_products (list[dict], optional): Products chosen in
              the product selector to include in Section 1.
        output_dir: Root output directory; summaries/ subdirectory is created.

    Returns:
        Absolute path to the saved .docx file.

    Raises:
        OSError: If the output directory cannot be created or the file cannot
            be written.
    """
    base_url = config.get("base_url", OLLAMA_BASE_URL)
    selected_products: list[dict[str, Any]] = config.get("selected_products", [])

    summaries_dir = Path(output_dir) / "summaries"
    try:
        summaries_dir.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        logger.error("Cannot create summaries directory '%s': %s", summaries_dir, exc)
        raise

    rfp_name = analysis.get("_rfp_filename", "rfp")
    stem = Path(rfp_name).stem
    customer_name = stem.replace("_", " ").replace("-", " ").title()
    date_str = datetime.now().strftime("%Y%m%d")
    display_date = datetime.now().strftime("%d %B %Y")

    filename = f"{stem[:30]}_Summary_{date_str}.docx"
    file_path = summaries_dir / filename

    logger.info("Generating customer summary document: %s", file_path)

    # ── Build compact context for the single call ─────────────────────────────
    exec_summary = (analysis.get("executive_summary") or "")[:300]
    industry = analysis.get("industry_vertical", "Unknown")
    product_names = (
        ", ".join(p.get("name", "") for p in selected_products)
        if selected_products
        else "TBC — to be selected"
    )
    requirements = analysis.get("requirements", [])
    mandatory = [r for r in requirements if str(r.get("type", "")).upper() == "MANDATORY"]
    top_reqs_lines = "\n".join(
        f"- {(r.get('requirement') or r.get('text') or r.get('description', ''))[:100]}"
        for r in mandatory[:8]
    )
    criteria = analysis.get("evaluation_criteria", [])
    criteria_lines = "\n".join(
        f"- {c.get('criterion', c.get('name', ''))} ({c.get('weight', 'no weight')})"
        for c in criteria[:10]
    )

    prompt = (
        "Generate a professional customer summary document with these 8 sections.\n"
        "Return XML only — no other text:\n\n"
        "<section1_objectives>Customer objectives content here</section1_objectives>\n"
        "<section2_use_cases>Use cases content here</section2_use_cases>\n"
        "<section3_technical>Technical requirements summary here</section3_technical>\n"
        "<section4_commercial>Commercial requirements here</section4_commercial>\n"
        "<section5_evaluation>Evaluation criteria content here</section5_evaluation>\n"
        "<section6_risks>Key risks and gaps here</section6_risks>\n"
        "<section7_questions>Clarification questions here</section7_questions>\n"
        "<section8_win_themes>Win themes here</section8_win_themes>\n\n"
        f"Context:\n"
        f"RFP Summary: {exec_summary}\n"
        f"Industry: {industry}\n"
        f"Products selected: {product_names}\n"
        f"Top requirements:\n{top_reqs_lines}\n"
        f"Evaluation criteria:\n{criteria_lines}"
    )

    # ── Single API call ───────────────────────────────────────────────────────
    payload: dict[str, Any] = {
        "model": SUMMARY_MODEL,
        "prompt": prompt,
        "stream": False,
        "options": {"num_predict": 2500, "temperature": 0.3, "top_p": 0.9},
    }
    logger.info("Calling Ollama '%s' for all summary sections (single call)", SUMMARY_MODEL)
    try:
        resp = requests.post(f"{base_url}/api/generate", json=payload, timeout=300)
        resp.raise_for_status()
        raw = resp.json().get("response", "")
    except requests.exceptions.RequestException as exc:
        logger.error("Single Ollama call failed: %s. Sections will use fallback text.", exc)
        raw = ""

    sections = _parse_xml_sections(raw)

    # ── Build Word document ───────────────────────────────────────────────────
    doc = Document()
    _add_page_numbers(doc)
    _add_title_page(doc, customer_name, customer_name, display_date)

    logger.info("Section 1: RFP Overview")
    _section1_rfp_overview(doc, analysis, selected_products)

    logger.info("Section 2: Customer Objectives")
    _section2_customer_objectives(doc, sections["section1_objectives"])

    logger.info("Section 3: Identified Use Cases")
    _section3_use_cases(doc, sections["section2_use_cases"])

    logger.info("Section 4: Technical Requirements Summary")
    _section4_technical_requirements(doc, analysis, sections["section3_technical"])

    logger.info("Section 5: Commercial Requirements")
    _section5_commercial_requirements(doc, sections["section4_commercial"])

    logger.info("Section 6: Evaluation Criteria")
    _section6_evaluation_criteria(doc, analysis, sections["section5_evaluation"])

    logger.info("Section 7: Key Risks and Gaps")
    _section7_risks_and_gaps(doc, sections["section6_risks"])

    logger.info("Section 8: Clarification Questions")
    _section8_clarification_questions(doc, sections["section7_questions"])

    logger.info("Section 9: Win Themes")
    _section9_win_themes(doc, sections["section8_win_themes"])

    try:
        doc.save(str(file_path))
    except OSError as exc:
        logger.error("Failed to save document '%s': %s", file_path, exc)
        raise

    resolved = str(file_path.resolve())
    logger.info("Customer summary saved: %s", resolved)
    return resolved
